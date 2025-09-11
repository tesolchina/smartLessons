#!/usr/bin/env python3
"""
Generic Document Analysis Framework
Reusable tool for analyzing Word documents and creating connections between multiple documents
"""

import docx
import re
import json
from pathlib import Path
from collections import defaultdict
from typing import Dict, List, Tuple, Any
import argparse

class GenericDocumentAnalyzer:
    """Generic framework for document analysis and comparison"""
    
    def __init__(self, config_file: str = None):
        """Initialize with optional configuration file"""
        self.config = self.load_config(config_file) if config_file else self.default_config()
    
    def default_config(self) -> Dict[str, Any]:
        """Default configuration for analysis"""
        return {
            'keywords': {
                'ai_technology': ['AI', 'artificial intelligence', 'AI-assisted', 'AI-based', 
                                'AI competency', 'AI literacy', 'AI training', 'AI tools',
                                'AI integration', 'AI workshops', 'AI Task Force',
                                'human-AI interaction', 'machine learning', 'automation'],
                'assessment': ['assessment', 'evaluation', 'grading', 'rubric', 'criteria',
                             'testing', 'examination', 'quality assurance', 'benchmarking'],
                'curriculum': ['curriculum', 'course design', 'syllabus', 'OBTL', 'PILO', 'CILO',
                              'learning outcomes', 'pedagogy', 'teaching methods'],
                'professional_development': ['professional development', 'training', 'workshop', 
                                           'conference', 'seminar', 'capacity building'],
                'internationalization': ['international', 'exchange', 'global citizenship', 
                                       'intercultural', 'multilingual', 'cross-cultural'],
                'community_engagement': ['community', 'public discourse', 'social impact', 
                                       'engagement', 'outreach', 'service'],
                'technology': ['technology', 'digital', 'e-tools', 'innovation', 'platform',
                             'online', 'virtual', 'multimedia']
            },
            'priority_weights': {
                'ai_technology': 5,
                'assessment': 4,
                'curriculum': 4,
                'professional_development': 3,
                'internationalization': 3,
                'community_engagement': 3,
                'technology': 4
            },
            'output_formats': ['json', 'markdown', 'html']
        }
    
    def load_config(self, config_file: str) -> Dict[str, Any]:
        """Load configuration from JSON file"""
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"Config file {config_file} not found. Using default configuration.")
            return self.default_config()
    
    def extract_document_content(self, doc_path: str) -> Dict[str, Any]:
        """Extract comprehensive content from Word document"""
        try:
            doc = docx.Document(doc_path)
            content = {
                'metadata': {
                    'filename': Path(doc_path).name,
                    'path': doc_path,
                    'paragraph_count': 0,
                    'table_count': 0
                },
                'paragraphs': [],
                'tables': [],
                'headers': [],
                'sections': []
            }
            
            # Extract paragraphs with metadata
            for para_idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    content['paragraphs'].append({
                        'index': para_idx,
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else 'Normal',
                        'is_heading': 'Heading' in (para.style.name if para.style else '')
                    })
            
            # Extract tables with structure
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    'table_id': table_idx + 1,
                    'rows': [],
                    'headers': []
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    
                    if row_idx == 0:
                        table_data['headers'] = row_data
                    table_data['rows'].append(row_data)
                
                content['tables'].append(table_data)
            
            # Extract headers and sections
            for para in content['paragraphs']:
                if para['is_heading']:
                    content['headers'].append({
                        'level': self.get_heading_level(para['style']),
                        'text': para['text'],
                        'index': para['index']
                    })
            
            content['metadata']['paragraph_count'] = len(content['paragraphs'])
            content['metadata']['table_count'] = len(content['tables'])
            
            return content
            
        except Exception as e:
            print(f"Error extracting {doc_path}: {e}")
            return None
    
    def get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        if 'Heading' in style_name:
            match = re.search(r'Heading (\d+)', style_name)
            return int(match.group(1)) if match else 1
        return 0
    
    def analyze_keywords(self, content: Dict[str, Any], custom_keywords: Dict[str, List[str]] = None) -> Dict[str, Any]:
        """Analyze keyword occurrences in document"""
        keywords = custom_keywords or self.config['keywords']
        analysis = {
            'total_mentions': 0,
            'categories': {},
            'detailed_mentions': {},
            'context_excerpts': {}
        }
        
        # Combine all text for analysis
        all_text = []
        text_sources = []
        
        # Add paragraphs
        for para in content['paragraphs']:
            all_text.append(para['text'])
            text_sources.append(('paragraph', para['index']))
        
        # Add table content
        for table in content['tables']:
            for row in table['rows']:
                text = ' | '.join(row)
                if text.strip():
                    all_text.append(text)
                    text_sources.append(('table', table['table_id']))
        
        # Analyze each category
        for category, category_keywords in keywords.items():
            analysis['categories'][category] = 0
            analysis['detailed_mentions'][category] = []
            analysis['context_excerpts'][category] = []
            
            for text_idx, text in enumerate(all_text):
                for keyword in category_keywords:
                    if keyword.lower() in text.lower():
                        analysis['total_mentions'] += 1
                        analysis['categories'][category] += 1
                        
                        mention = {
                            'keyword': keyword,
                            'context': text,
                            'source_type': text_sources[text_idx][0],
                            'source_id': text_sources[text_idx][1],
                            'context_length': len(text)
                        }
                        
                        analysis['detailed_mentions'][category].append(mention)
                        
                        # Extract context excerpt (150 chars around keyword)
                        keyword_pos = text.lower().find(keyword.lower())
                        start = max(0, keyword_pos - 75)
                        end = min(len(text), keyword_pos + 75)
                        excerpt = text[start:end]
                        
                        if excerpt not in analysis['context_excerpts'][category]:
                            analysis['context_excerpts'][category].append(excerpt)
        
        return analysis
    
    def find_document_connections(self, doc1_content: Dict[str, Any], doc2_content: Dict[str, Any], 
                                analysis1: Dict[str, Any], analysis2: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Find thematic connections between two documents"""
        connections = []
        
        for category in self.config['keywords'].keys():
            if (analysis1['categories'].get(category, 0) > 0 and 
                analysis2['categories'].get(category, 0) > 0):
                
                connection_strength = min(
                    analysis1['categories'][category],
                    analysis2['categories'][category]
                )
                
                connection = {
                    'category': category,
                    'strength': connection_strength,
                    'priority': self.config['priority_weights'].get(category, 1),
                    'doc1_mentions': analysis1['categories'][category],
                    'doc2_mentions': analysis2['categories'][category],
                    'doc1_excerpts': analysis1['context_excerpts'][category][:3],
                    'doc2_excerpts': analysis2['context_excerpts'][category][:3],
                    'enhancement_score': connection_strength * self.config['priority_weights'].get(category, 1)
                }
                
                connections.append(connection)
        
        return sorted(connections, key=lambda x: x['enhancement_score'], reverse=True)
    
    def generate_enhanced_markdown(self, content: Dict[str, Any], analysis: Dict[str, Any], 
                                 connections: List[Dict[str, Any]] = None) -> str:
        """Generate enhanced markdown with excerpts and analysis"""
        
        doc_name = content['metadata']['filename']
        
        markdown = f"""# Enhanced Analysis: {doc_name}

## Document Overview
- **Paragraphs:** {content['metadata']['paragraph_count']}
- **Tables:** {content['metadata']['table_count']}  
- **Total Keyword Mentions:** {analysis['total_mentions']}

## Keyword Analysis by Category

"""
        
        for category, count in analysis['categories'].items():
            if count > 0:
                priority = self.config['priority_weights'].get(category, 1)
                priority_indicator = "🔴" if priority >= 4 else "🟡" if priority >= 3 else "🟢"
                
                markdown += f"""### {category.replace('_', ' ').title()} {priority_indicator}
**Mentions:** {count} | **Priority:** {priority}/5

#### Key Excerpts:
"""
                # Add top 3 context excerpts
                for excerpt in analysis['context_excerpts'][category][:3]:
                    markdown += f"- *...{excerpt}...*\n"
                
                markdown += "\n"
        
        if connections:
            markdown += """
## Document Connections

"""
            for conn in connections[:5]:  # Top 5 connections
                markdown += f"""### {conn['category'].replace('_', ' ').title()}
**Strength:** {conn['strength']} | **Enhancement Score:** {conn['enhancement_score']}

#### Supporting Evidence:
"""
                for excerpt in conn['doc1_excerpts']:
                    markdown += f"- *{excerpt}*\n"
        
        return markdown
    
    def save_results(self, results: Dict[str, Any], output_path: str, format_type: str = 'json'):
        """Save results in specified format"""
        output_file = Path(output_path)
        
        if format_type == 'json':
            with open(output_file.with_suffix('.json'), 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
        
        elif format_type == 'markdown':
            with open(output_file.with_suffix('.md'), 'w', encoding='utf-8') as f:
                f.write(results)
        
        print(f"✅ Results saved to: {output_file.with_suffix('.' + format_type)}")

def main():
    """Main function with CLI interface"""
    parser = argparse.ArgumentParser(description='Generic Document Analysis Framework')
    parser.add_argument('documents', nargs='+', help='Documents to analyze')
    parser.add_argument('--config', help='Configuration file path')
    parser.add_argument('--output', default='analysis_results', help='Output file prefix')
    parser.add_argument('--format', choices=['json', 'markdown'], default='json', help='Output format')
    parser.add_argument('--keywords', help='Custom keywords file (JSON)')
    
    args = parser.parse_args()
    
    analyzer = GenericDocumentAnalyzer(args.config)
    
    # Analyze documents
    documents_data = []
    for doc_path in args.documents:
        print(f"📄 Analyzing: {doc_path}")
        content = analyzer.extract_document_content(doc_path)
        if content:
            analysis = analyzer.analyze_keywords(content)
            documents_data.append({
                'path': doc_path,
                'content': content,
                'analysis': analysis
            })
    
    # Find connections if multiple documents
    connections = []
    if len(documents_data) >= 2:
        print("🔗 Finding document connections...")
        connections = analyzer.find_document_connections(
            documents_data[0]['content'],
            documents_data[1]['content'],
            documents_data[0]['analysis'],
            documents_data[1]['analysis']
        )
    
    # Generate results
    results = {
        'documents': documents_data,
        'connections': connections,
        'analysis_summary': {
            'total_documents': len(documents_data),
            'total_connections': len(connections)
        }
    }
    
    analyzer.save_results(results, args.output, args.format)
    print(f"🎯 Analysis complete!")

if __name__ == "__main__":
    main()
