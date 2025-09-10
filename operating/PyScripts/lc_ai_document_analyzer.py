#!/usr/bin/env python3
"""
AI Analysis and Document Connection Tool for LC Admin
Analyzes DAA Report for AI mentions and creates connections with Language Education Paper
"""

import docx
import re
import json
from pathlib import Path
from collections import defaultdict

class LCDocumentAnalyzer:
    def __init__(self):
        self.ai_keywords = [
            'AI', 'artificial intelligence', 'AI-assisted', 'AI-based', 
            'AI competency', 'AI literacy', 'AI training', 'AI tools',
            'AI integration', 'AI workshops', 'AI Task Force',
            'human-AI interaction', 'technology', 'digital', 'e-tools'
        ]
        
    def extract_docx_content(self, doc_path):
        """Extract content from Word document"""
        try:
            doc = docx.Document(doc_path)
            content = {
                'paragraphs': [],
                'tables': []
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append(para.text.strip())
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                content['tables'].append({
                    'table_id': table_idx + 1,
                    'data': table_data
                })
            
            return content
            
        except Exception as e:
            print(f"Error extracting {doc_path}: {e}")
            return None
    
    def find_ai_mentions(self, text_content):
        """Find all AI-related mentions in text"""
        ai_mentions = []
        
        # Combine all text
        all_text = []
        if 'paragraphs' in text_content:
            all_text.extend(text_content['paragraphs'])
        
        if 'tables' in text_content:
            for table in text_content['tables']:
                for row in table['data']:
                    all_text.extend([cell for cell in row if cell])
        
        # Search for AI keywords
        for line_idx, line in enumerate(all_text):
            for keyword in self.ai_keywords:
                if keyword.lower() in line.lower():
                    ai_mentions.append({
                        'line_number': line_idx + 1,
                        'keyword': keyword,
                        'context': line,
                        'source': 'paragraph' if line_idx < len(text_content.get('paragraphs', [])) else 'table'
                    })
        
        return ai_mentions
    
    def categorize_ai_mentions(self, ai_mentions):
        """Categorize AI mentions by section and type"""
        categories = {
            'ai_competency_development': [],
            'ai_assisted_learning': [],
            'ai_assessment': [],
            'ai_training_development': [],
            'ai_tools_integration': [],
            'ai_policy_strategy': []
        }
        
        for mention in ai_mentions:
            context_lower = mention['context'].lower()
            
            if any(word in context_lower for word in ['competency', 'competence', 'skill']):
                categories['ai_competency_development'].append(mention)
            elif any(word in context_lower for word in ['assessment', 'evaluate', 'grading']):
                categories['ai_assessment'].append(mention)
            elif any(word in context_lower for word in ['training', 'workshop', 'development']):
                categories['ai_training_development'].append(mention)
            elif any(word in context_lower for word in ['tool', 'platform', 'technology']):
                categories['ai_tools_integration'].append(mention)
            elif any(word in context_lower for word in ['policy', 'strategy', 'task force']):
                categories['ai_policy_strategy'].append(mention)
            else:
                categories['ai_assisted_learning'].append(mention)
        
        return categories
    
    def find_thematic_connections(self, daa_content, lang_ed_content):
        """Find thematic connections between documents"""
        connections = []
        
        # Define connection themes
        themes = {
            'multilingual_education': ['multilingual', 'foreign language', 'trilingual', 'language learning'],
            'assessment_quality': ['assessment', 'quality assurance', 'evaluation', 'standard'],
            'curriculum_development': ['curriculum', 'course design', 'syllabus', 'OBTL'],
            'professional_development': ['professional development', 'training', 'workshop', 'conference'],
            'international_programs': ['international', 'exchange', 'global citizenship', 'intercultural'],
            'community_engagement': ['community', 'public discourse', 'social impact', 'engagement'],
            'technology_innovation': ['AI', 'technology', 'digital', 'innovation', 'e-tools']
        }
        
        # Search for theme connections
        daa_text = ' '.join(daa_content.get('paragraphs', []))
        lang_ed_text = ' '.join(lang_ed_content.get('paragraphs', []))
        
        for theme, keywords in themes.items():
            daa_matches = []
            lang_ed_matches = []
            
            for keyword in keywords:
                if keyword.lower() in daa_text.lower():
                    daa_matches.append(keyword)
                if keyword.lower() in lang_ed_text.lower():
                    lang_ed_matches.append(keyword)
            
            if daa_matches and lang_ed_matches:
                connections.append({
                    'theme': theme,
                    'daa_keywords': daa_matches,
                    'lang_ed_keywords': lang_ed_matches,
                    'connection_strength': len(set(daa_matches) & set(lang_ed_matches))
                })
        
        return sorted(connections, key=lambda x: x['connection_strength'], reverse=True)

def main():
    analyzer = LCDocumentAnalyzer()
    
    # Document paths
    daa_path = "/Users/simonwang/Documents/Usage/VibeCoding/DailyAssistant/projects/LC_admin/PMC/LC DAA Report_response_20251009.docx"
    lang_ed_path = "/Users/simonwang/Documents/Usage/VibeCoding/DailyAssistant/projects/LC_admin/Departmental_meeting/Language Education Paper Clean 20250409.docx"
    
    print("🔍 LC AI Analysis and Document Connection Tool")
    print("=" * 50)
    
    # Extract DAA Report content
    print("📄 Extracting DAA Report...")
    daa_content = analyzer.extract_docx_content(daa_path)
    if not daa_content:
        print("❌ Failed to extract DAA Report")
        return
    
    # Extract Language Education Paper content
    print("📄 Extracting Language Education Paper...")
    lang_ed_content = analyzer.extract_docx_content(lang_ed_path)
    if not lang_ed_content:
        print("❌ Failed to extract Language Education Paper")
        return
    
    # Find AI mentions in DAA Report
    print("\n🤖 Analyzing AI mentions in DAA Report...")
    ai_mentions = analyzer.find_ai_mentions(daa_content)
    ai_categories = analyzer.categorize_ai_mentions(ai_mentions)
    
    print(f"✅ Found {len(ai_mentions)} AI-related mentions")
    
    # Find thematic connections
    print("\n🔗 Finding thematic connections between documents...")
    connections = analyzer.find_thematic_connections(daa_content, lang_ed_content)
    
    # Generate analysis results
    results = {
        'ai_analysis': {
            'total_mentions': len(ai_mentions),
            'categories': {k: len(v) for k, v in ai_categories.items()},
            'detailed_mentions': ai_categories
        },
        'document_connections': connections,
        'recommendations': []
    }
    
    # Save results
    with open('ai_analysis_results.json', 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    
    print(f"✅ Analysis complete - {len(connections)} thematic connections found")
    print("📊 Results saved to ai_analysis_results.json")
    
    return results

if __name__ == "__main__":
    results = main()
