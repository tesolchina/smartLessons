#!/usr/bin/env python3
"""
Convert all .docx files in the project to markdown format
"""

import os
from docx import Document
import re

def docx_to_markdown(docx_path, md_path):
    """Convert a .docx file to markdown format"""
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                markdown_content.append("")
                continue
                
            # Check paragraph style for headers
            style_name = ""
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name.lower()
            
            if 'heading 1' in style_name or 'title' in style_name:
                markdown_content.append(f"# {text}")
            elif 'heading 2' in style_name:
                markdown_content.append(f"## {text}")
            elif 'heading 3' in style_name:
                markdown_content.append(f"### {text}")
            elif 'heading 4' in style_name:
                markdown_content.append(f"#### {text}")
            else:
                # Regular paragraph
                markdown_content.append(text)
        
        # Handle tables
        for table in doc.tables:
            markdown_content.append("\n")
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_data.append(cell_text)
                
                if i == 0:  # Header row
                    markdown_content.append("| " + " | ".join(row_data) + " |")
                    markdown_content.append("| " + " | ".join(["---"] * len(row_data)) + " |")
                else:
                    markdown_content.append("| " + " | ".join(row_data) + " |")
            markdown_content.append("\n")
        
        # Write to markdown file
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_content))
        
        print(f"✅ Converted: {os.path.basename(docx_path)} → {os.path.basename(md_path)}")
        return True
        
    except Exception as e:
        print(f"❌ Error converting {docx_path}: {str(e)}")
        return False

def find_and_convert_docx_files(root_dir):
    """Find all .docx files and convert them to markdown"""
    converted_files = []
    
    for root, dirs, files in os.walk(root_dir):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~'):  # Skip temp files
                docx_path = os.path.join(root, file)
                md_filename = file.replace('.docx', '.md')
                md_path = os.path.join(root, md_filename)
                
                print(f"Converting: {docx_path}")
                if docx_to_markdown(docx_path, md_path):
                    converted_files.append((docx_path, md_path))
    
    return converted_files

if __name__ == "__main__":
    project_root = "/Users/simonwang/Documents/Usage/VibeCodingMac/DailyAssistant/projects/Donald"
    print("🔄 Starting conversion of .docx files to markdown...")
    
    converted = find_and_convert_docx_files(project_root)
    
    print(f"\n📊 Conversion Summary:")
    print(f"Total files converted: {len(converted)}")
    
    for docx_path, md_path in converted:
        print(f"  • {os.path.basename(docx_path)} → {os.path.basename(md_path)}")
    
    print("\n✅ Conversion complete!")