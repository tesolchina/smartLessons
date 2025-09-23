#!/usr/bin/env python3
"""
TKP Foundation Report Form Filler
================================
Automated form filling for the TKP Foundation Word document using python-docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from datetime import datetime

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class TKPFormFiller:
    def __init__(self):
        self.reports_dir = Path("/Users/simonwang/Documents/Usage/VibeCodingMac/DailyAssistant/projects/LANG2077/CourseDocs/Reports")
        self.tkp_dir = self.reports_dir / "TKPfoundation"
        
        # Template and output paths
        self.template_file = self.tkp_dir / "田家炳基金會 青年品格培育計劃報告表_2024-25_template_LANG2077.docx"
        self.output_file = self.tkp_dir / "田家炳基金會_LANG2077_Report_Completed.docx"
        
        # Data structure to hold form responses
        self.form_data = self.initialize_form_data()
    
    def initialize_form_data(self):
        """Initialize form data structure with default values."""
        return {
            # Instructor Information
            "instructor_name": "Prof. Simon Wang",
            "department": "Language Centre 語文中心",
            "email": "simonwang@hkbu.edu.hk",
            "extension": "3943",
            
            # Course Information
            "course_code": "LANG2077",
            "course_title": "Language Skills for Human-AI Partnership: Customizing Chatbots to Empower Communities",
            "semester": "May – August 2025",
            "total_students": "31",
            "project_assistants": "2",
            
            # Project Outcomes - to be filled with bullet points
            "overall_impact": "",
            "academic_application": "",
            "student_changes": "",
            "cultural_understanding": "",
            "cultural_recognition": "",
            "challenges_solutions": "",
            
            # Character Development (checkboxes)
            "character_traits": {
                "national_identity": True,
                "responsibility": True,
                "citizenship": True,
                "commitment": True,
                "integrity": True,
                "empathy": True,
                "teamwork": True,
                "creativity": True,
                "problem_solving": True,
                "perseverance": True,
                "respect": True,
                "caring": True
            },
            
            # SDGs (checkboxes)
            "sdgs": {
                "sdg04": True,  # Quality Education
                "sdg05": True,  # Gender Equality
                "sdg10": True,  # Reduced Inequalities
                "sdg16": True,  # Peace, Justice and Strong Institutions
                "sdg17": True   # Partnerships for the Goals
            },
            
            # TKP Goals (checkboxes)
            "tkp_goals": {
                "chinese_culture_role": True,
                "humanistic_qualities": True,
                "national_responsibility": True,
                "multiple_citizenship": True
            },
            
            # Financial Information
            "course_subsidy": "$30,000",
            "student_subsidy": "$15,500",
            "travel_expenses_students": "$144,522",
            "travel_expenses_instructors": "$10,000",
            "closing_balance": "$200,022",
            
            # Future Plans
            "future_interest": True,
            "suggestions": "",
            "other_comments": "",
            "scheme_elaboration": ""
        }
    
    def update_form_data(self, bullet_points_dict):
        """Update form data with provided bullet points."""
        if "overall_impact" in bullet_points_dict:
            self.form_data["overall_impact"] = self.format_bullet_points(bullet_points_dict["overall_impact"])
        
        if "academic_application" in bullet_points_dict:
            self.form_data["academic_application"] = self.format_bullet_points(bullet_points_dict["academic_application"])
        
        if "student_changes" in bullet_points_dict:
            self.form_data["student_changes"] = self.format_bullet_points(bullet_points_dict["student_changes"])
        
        if "cultural_understanding" in bullet_points_dict:
            self.form_data["cultural_understanding"] = self.format_bullet_points(bullet_points_dict["cultural_understanding"])
        
        if "cultural_recognition" in bullet_points_dict:
            self.form_data["cultural_recognition"] = self.format_bullet_points(bullet_points_dict["cultural_recognition"])
        
        if "challenges_solutions" in bullet_points_dict:
            self.form_data["challenges_solutions"] = self.format_bullet_points(bullet_points_dict["challenges_solutions"])
        
        if "suggestions" in bullet_points_dict:
            self.form_data["suggestions"] = self.format_bullet_points(bullet_points_dict["suggestions"])
        
        if "other_comments" in bullet_points_dict:
            self.form_data["other_comments"] = self.format_bullet_points(bullet_points_dict["other_comments"])
        
        if "scheme_elaboration" in bullet_points_dict:
            self.form_data["scheme_elaboration"] = self.format_bullet_points(bullet_points_dict["scheme_elaboration"])
    
    def format_bullet_points(self, bullet_points):
        """Format bullet points into readable text."""
        if isinstance(bullet_points, list):
            return "\n".join([f"• {point}" for point in bullet_points])
        elif isinstance(bullet_points, str):
            return bullet_points
        else:
            return str(bullet_points)
    
    def fill_form(self, bullet_points_dict=None):
        """Fill the Word document form with data."""
        logger.info("Starting TKP Foundation form filling...")
        
        # Update form data if bullet points provided
        if bullet_points_dict:
            self.update_form_data(bullet_points_dict)
        
        try:
            # Load the template document
            doc = Document(self.template_file)
            logger.info(f"Loaded template: {self.template_file.name}")
            
            # Fill basic information by text replacement
            self.replace_text_in_document(doc)
            
            # Handle checkboxes and special formatting
            self.handle_checkboxes(doc)
            
            # Save the completed document
            doc.save(self.output_file)
            logger.info(f"✅ Form completed and saved: {self.output_file.name}")
            
            return str(self.output_file)
            
        except Exception as e:
            logger.error(f"❌ Error filling form: {str(e)}")
            raise
    
    def replace_text_in_document(self, doc):
        """Replace placeholder text in the document."""
        replacements = {
            # Basic Information Placeholders
            "[INSTRUCTOR_NAME]": self.form_data["instructor_name"],
            "[DEPARTMENT]": self.form_data["department"],
            "[EMAIL]": self.form_data["email"],
            "[EXTENSION]": self.form_data["extension"],
            "[COURSE_CODE]": self.form_data["course_code"],
            "[COURSE_TITLE]": self.form_data["course_title"],
            "[SEMESTER]": self.form_data["semester"],
            "[TOTAL_STUDENTS]": self.form_data["total_students"],
            "[PROJECT_ASSISTANTS]": self.form_data["project_assistants"],
            
            # Content Placeholders
            "[OVERALL_IMPACT]": self.form_data["overall_impact"],
            "[ACADEMIC_APPLICATION]": self.form_data["academic_application"],
            "[STUDENT_CHANGES]": self.form_data["student_changes"],
            "[CULTURAL_UNDERSTANDING]": self.form_data["cultural_understanding"],
            "[CULTURAL_RECOGNITION]": self.form_data["cultural_recognition"],
            "[CHALLENGES_SOLUTIONS]": self.form_data["challenges_solutions"],
            "[SUGGESTIONS]": self.form_data["suggestions"],
            "[OTHER_COMMENTS]": self.form_data["other_comments"],
            "[SCHEME_ELABORATION]": self.form_data["scheme_elaboration"],
            
            # Financial Information
            "[COURSE_SUBSIDY]": self.form_data["course_subsidy"],
            "[STUDENT_SUBSIDY]": self.form_data["student_subsidy"],
            "[TRAVEL_STUDENTS]": self.form_data["travel_expenses_students"],
            "[TRAVEL_INSTRUCTORS]": self.form_data["travel_expenses_instructors"],
            "[CLOSING_BALANCE]": self.form_data["closing_balance"],
        }
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, replacement in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, replacement)
    
    def handle_checkboxes(self, doc):
        """Handle checkbox selections in the document."""
        # This is a simplified approach - in practice, you might need to
        # identify specific checkbox patterns and mark them
        
        checkbox_replacements = {}
        
        # Character traits checkboxes
        if self.form_data["character_traits"]["national_identity"]:
            checkbox_replacements["☐ National Identity"] = "☑ National Identity"
        if self.form_data["character_traits"]["responsibility"]:
            checkbox_replacements["☐ Responsibility"] = "☑ Responsibility"
        if self.form_data["character_traits"]["teamwork"]:
            checkbox_replacements["☐ Teamwork"] = "☑ Teamwork"
        
        # SDG checkboxes
        if self.form_data["sdgs"]["sdg04"]:
            checkbox_replacements["☐ SDG 04"] = "☑ SDG 04"
        
        # Apply checkbox replacements
        for paragraph in doc.paragraphs:
            for unchecked, checked in checkbox_replacements.items():
                if unchecked in paragraph.text:
                    paragraph.text = paragraph.text.replace(unchecked, checked)
    
    def test_form_filling(self):
        """Test the form filling with sample data."""
        logger.info("Testing form filling with sample data...")
        
        sample_bullet_points = {
            "overall_impact": [
                "Students developed strong ethical reasoning skills through AI education",
                "Enhanced cultural awareness through service-learning in rural China",
                "Promoted cross-cultural understanding and respect"
            ],
            "academic_application": [
                "Applied programming skills to create educational games",
                "Used language and communication skills for teaching",
                "Implemented pedagogical theories in practical teaching scenarios"
            ],
            "student_changes": [
                "Increased empathy and cultural sensitivity",
                "Improved leadership and teamwork abilities",
                "Enhanced problem-solving and adaptability skills"
            ],
            "challenges_solutions": [
                "Language barriers - resolved through peer translation and visual aids",
                "Technology limitations - adapted materials for offline use",
                "Cultural differences - addressed through cultural orientation sessions"
            ]
        }
        
        return self.fill_form(sample_bullet_points)
    
    def create_form_template_with_placeholders(self):
        """Create a version of the form with placeholders for easier filling."""
        logger.info("Creating form template with placeholders...")
        
        try:
            # Load the original template
            doc = Document(self.template_file)
            
            # Add placeholders in strategic locations
            placeholder_doc = self.add_placeholders_to_document(doc)
            
            # Save template with placeholders
            placeholder_file = self.tkp_dir / "TKP_Foundation_Template_With_Placeholders.docx"
            doc.save(placeholder_file)
            
            logger.info(f"✅ Template with placeholders saved: {placeholder_file.name}")
            return str(placeholder_file)
            
        except Exception as e:
            logger.error(f"❌ Error creating placeholder template: {str(e)}")
            raise
    
    def add_placeholders_to_document(self, doc):
        """Add placeholders to the document for easier automated filling."""
        # This would involve identifying specific locations in the document
        # and inserting placeholder text that can be easily replaced later
        
        # For now, return the document as-is
        # In practice, you'd analyze the document structure and add placeholders
        return doc

# Test and usage functions
def test_tkp_form_filler():
    """Test the TKP form filler functionality."""
    filler = TKPFormFiller()
    
    # Test with sample data
    output_file = filler.test_form_filling()
    print(f"✅ Test completed. Output file: {output_file}")
    
    return output_file

def create_placeholder_template():
    """Create a template with placeholders for easier form filling."""
    filler = TKPFormFiller()
    template_file = filler.create_form_template_with_placeholders()
    print(f"✅ Placeholder template created: {template_file}")
    return template_file

if __name__ == "__main__":
    print("🔧 TKP Foundation Form Filler")
    print("=" * 40)
    
    # Create form filler instance
    filler = TKPFormFiller()
    
    # Option 1: Test with sample data
    print("\n1. Testing with sample data...")
    try:
        test_output = test_tkp_form_filler()
        print(f"✅ Sample form created: {test_output}")
    except Exception as e:
        print(f"❌ Test failed: {e}")
    
    # Option 2: Create placeholder template
    print("\n2. Creating placeholder template...")
    try:
        placeholder_template = create_placeholder_template()
        print(f"✅ Placeholder template: {placeholder_template}")
    except Exception as e:
        print(f"❌ Template creation failed: {e}")
    
    print("\n🎯 Next Steps:")
    print("1. Provide bullet points for each section")
    print("2. Run filler.fill_form(bullet_points_dict) to generate completed form")
    print("3. Review and submit the completed document")