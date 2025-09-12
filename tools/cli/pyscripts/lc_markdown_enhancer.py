#!/usr/bin/env python3
"""
LC Document Enhancement Tool
Specifically designed for enhancing LC admin markdown files with excerpts
"""

import docx
import re
import json
from pathlib import Path
from typing import Dict, List

def extract_word_content(doc_path: str) -> Dict[str, List[str]]:
    """Extract content from Word document"""
    try:
        doc = docx.Document(doc_path)
        content = {'paragraphs': [], 'ai_mentions': [], 'all_excerpts': []}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 50:  # Filter out short paragraphs
                content['paragraphs'].append(text)
                
                # Identify AI-related content
                ai_keywords = ['AI', 'artificial intelligence', 'AI-assisted', 'AI competency', 
                              'AI literacy', 'AI training', 'AI tools', 'AI integration']
                
                if any(keyword.lower() in text.lower() for keyword in ai_keywords):
                    content['ai_mentions'].append(text)
                
                content['all_excerpts'].append(text)
        
        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                if row_text.strip() and len(row_text) > 30:
                    content['all_excerpts'].append(row_text)
                    
                    ai_keywords = ['AI', 'artificial intelligence', 'AI-assisted', 'AI competency']
                    if any(keyword.lower() in row_text.lower() for keyword in ai_keywords):
                        content['ai_mentions'].append(row_text)
        
        return content
    
    except Exception as e:
        print(f"Error processing {doc_path}: {e}")
        return {'paragraphs': [], 'ai_mentions': [], 'all_excerpts': []}

def find_topic_excerpts(content: Dict[str, List[str]], topic_keywords: List[str], max_excerpts: int = 3) -> List[str]:
    """Find excerpts related to specific topics"""
    scored_excerpts = []
    
    for text in content['all_excerpts']:
        score = sum(1 for keyword in topic_keywords if keyword.lower() in text.lower())
        if score > 0:
            scored_excerpts.append({
                'text': text,
                'score': score,
                'length': len(text)
            })
    
    # Sort by score (relevance) and length
    scored_excerpts.sort(key=lambda x: (x['score'], min(x['length'], 300)), reverse=True)
    
    # Return top excerpts, truncated if too long
    results = []
    for excerpt in scored_excerpts[:max_excerpts]:
        text = excerpt['text']
        if len(text) > 250:
            text = text[:247] + "..."
        results.append(text)
    
    return results

def enhance_lc_markdown(markdown_path: str, daa_report_path: str, lang_ed_paper_path: str) -> str:
    """Enhance LC markdown file with excerpts from both source documents"""
    
    # Read the markdown file
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Extract content from source documents
    print(f"📄 Processing DAA Report...")
    daa_content = extract_word_content(daa_report_path)
    
    print(f"📄 Processing Language Education Paper...")
    lang_ed_content = extract_word_content(lang_ed_paper_path)
    
    # Define topic mappings for different sections
    topic_mappings = {
        'AI': ['AI', 'artificial intelligence', 'AI-assisted', 'AI competency', 'AI literacy', 'AI training', 'AI tools'],
        'Assessment': ['assessment', 'evaluation', 'rubric', 'criteria', 'testing', 'quality assurance'],
        'Curriculum': ['curriculum', 'OBTL', 'PILO', 'CILO', 'learning outcomes', 'course design'],
        'Multilingual': ['multilingual', 'English', 'Chinese', 'language learning', 'bilingual'],
        'Professional Development': ['professional development', 'training', 'workshop', 'capacity building'],
        'International': ['international', 'exchange', 'global citizenship', 'intercultural']
    }
    
    enhanced_content = content
    
    # Add comprehensive evidence section at the end
    evidence_section = f"\n\n## Supporting Evidence and Excerpts\n\n"
    evidence_section += "### From DAA Report\n\n"
    
    # Add AI-related excerpts
    if daa_content['ai_mentions']:
        evidence_section += "**AI and Technology Integration:**\n\n"
        for i, excerpt in enumerate(daa_content['ai_mentions'][:3], 1):
            if len(excerpt) > 200:
                excerpt = excerpt[:197] + "..."
            evidence_section += f"{i}. *\"{excerpt}\"*\n\n"
    
    # Add general strategic excerpts
    strategic_keywords = ['strategic', 'mission', 'vision', 'direction', 'planning', 'development']
    strategic_excerpts = find_topic_excerpts(daa_content, strategic_keywords, 2)
    if strategic_excerpts:
        evidence_section += "**Strategic Direction:**\n\n"
        for i, excerpt in enumerate(strategic_excerpts, 1):
            evidence_section += f"{i}. *\"{excerpt}\"*\n\n"
    
    evidence_section += "### From Language Education Paper\n\n"
    
    # Add AI-related excerpts from language education paper
    if lang_ed_content['ai_mentions']:
        evidence_section += "**AI in Language Education:**\n\n"
        for i, excerpt in enumerate(lang_ed_content['ai_mentions'][:3], 1):
            if len(excerpt) > 200:
                excerpt = excerpt[:197] + "..."
            evidence_section += f"{i}. *\"{excerpt}\"*\n\n"
    
    # Add curriculum and assessment excerpts
    for topic, keywords in [('Curriculum Design', topic_mappings['Curriculum']), 
                           ('Assessment Innovation', topic_mappings['Assessment'])]:
        topic_excerpts = find_topic_excerpts(lang_ed_content, keywords, 2)
        if topic_excerpts:
            evidence_section += f"**{topic}:**\n\n"
            for i, excerpt in enumerate(topic_excerpts, 1):
                evidence_section += f"{i}. *\"{excerpt}\"*\n\n"
    
    # Add cross-document connections
    evidence_section += "### Cross-Document Strategic Connections\n\n"
    
    # Find AI competency connections
    ai_daa = len(daa_content['ai_mentions'])
    ai_lang = len(lang_ed_content['ai_mentions'])
    
    if ai_daa > 0 and ai_lang > 0:
        evidence_section += f"**AI Integration Alignment:** Both documents emphasize AI competency development with {ai_daa} mentions in DAA Report and {ai_lang} mentions in Language Education Paper, indicating strong institutional commitment to AI integration.\n\n"
    
    # Find multilingual connections
    multilingual_excerpts_daa = find_topic_excerpts(daa_content, topic_mappings['Multilingual'], 1)
    multilingual_excerpts_lang = find_topic_excerpts(lang_ed_content, topic_mappings['Multilingual'], 1)
    
    if multilingual_excerpts_daa and multilingual_excerpts_lang:
        evidence_section += "**Multilingual Education Synergy:** Both documents support multilingual learning approaches:\n\n"
        evidence_section += f"- DAA Report: *\"{multilingual_excerpts_daa[0]}\"*\n\n"
        evidence_section += f"- Language Education Paper: *\"{multilingual_excerpts_lang[0]}\"*\n\n"
    
    enhanced_content = enhanced_content + evidence_section
    
    return enhanced_content

def main():
    """Enhance all LC markdown files with excerpts"""
    
    base_dir = Path("/Users/simonwang/Documents/Usage/VibeCoding/DailyAssistant/projects/LC_admin/PMC")
    daa_report = base_dir / "LC DAA Report_response_20251009.docx"
    lang_ed_paper = Path("/Users/simonwang/Documents/Usage/VibeCoding/DailyAssistant/projects/LC_admin/Departmental_meeting/Language Education Paper Clean 20250409.docx")
    
    markdown_files = [
        "01_scope_strategic_direction.md",
        "02_staffing.md", 
        "03_teaching_learning.md",
        "04_research_scholarly_activities.md",
        "05_community_impact_internationalisation.md"
    ]
    
    for md_file in markdown_files:
        md_path = base_dir / md_file
        if md_path.exists():
            print(f"🔧 Enhancing: {md_file}")
            
            enhanced_content = enhance_lc_markdown(str(md_path), str(daa_report), str(lang_ed_paper))
            
            # Save enhanced version
            output_path = base_dir / f"enhanced_{md_file}"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(enhanced_content)
            
            print(f"✅ Enhanced version saved: enhanced_{md_file}")
        else:
            print(f"⚠️  File not found: {md_file}")
    
    print("\n🎯 All files enhanced with supporting evidence and excerpts!")

if __name__ == "__main__":
    main()
