#!/usr/bin/env python3
"""
Universal Document Enhancement Framework
Configurable tool for enhancing any markdown files with excerpts from multiple source documents
"""

import docx
import json
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
import argparse
from dataclasses import dataclass, field

@dataclass
class EnhancementConfig:
    """Configuration class for document enhancement"""
    topic_mappings: Dict[str, List[str]] = field(default_factory=dict)
    max_excerpts_per_topic: int = 3
    excerpt_max_length: int = 250
    source_documents: List[str] = field(default_factory=list)
    output_format: str = 'markdown'
    cross_reference_analysis: bool = True
    
    @classmethod
    def from_json(cls, config_path: str) -> 'EnhancementConfig':
        """Load configuration from JSON file"""
        with open(config_path, 'r', encoding='utf-8') as f:
            config_data = json.load(f)
        
        return cls(**config_data)
    
    def save_template(self, path: str):
        """Save a template configuration file"""
        template = {
            "topic_mappings": {
                "AI_Technology": ["AI", "artificial intelligence", "machine learning", "automation"],
                "Assessment": ["assessment", "evaluation", "rubric", "testing", "quality assurance"],
                "Curriculum": ["curriculum", "course design", "learning outcomes", "pedagogy"],
                "Professional_Development": ["training", "workshop", "professional development"],
                "International": ["international", "global", "exchange", "intercultural"]
            },
            "max_excerpts_per_topic": 3,
            "excerpt_max_length": 250,
            "source_documents": ["document1.docx", "document2.docx"],
            "output_format": "markdown",
            "cross_reference_analysis": True
        }
        
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(template, f, indent=2, ensure_ascii=False)
        
        print(f"✅ Template configuration saved to: {path}")

class UniversalDocumentEnhancer:
    """Universal framework for enhancing documents with configurable topics and sources"""
    
    def __init__(self, config: EnhancementConfig):
        self.config = config
        self.source_contents: Dict[str, Dict[str, List[str]]] = {}
    
    def load_source_documents(self):
        """Load content from all source documents"""
        for doc_path in self.config.source_documents:
            if Path(doc_path).exists():
                print(f"📄 Loading: {Path(doc_path).name}")
                self.source_contents[doc_path] = self.extract_document_content(doc_path)
            else:
                print(f"⚠️  Document not found: {doc_path}")
    
    def extract_document_content(self, doc_path: str) -> Dict[str, List[str]]:
        """Extract structured content from Word document"""
        try:
            doc = docx.Document(doc_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'topic_excerpts': {},
                'all_text': []
            }
            
            # Initialize topic collections
            for topic in self.config.topic_mappings.keys():
                content['topic_excerpts'][topic] = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) > 30:  # Filter meaningful content
                    content['paragraphs'].append(text)
                    content['all_text'].append(text)
                    
                    # Categorize by topics
                    self.categorize_text(text, content['topic_excerpts'])
            
            # Process tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_rows.append(row_text)
                        content['all_text'].append(row_text)
                        self.categorize_text(row_text, content['topic_excerpts'])
                
                if table_rows:
                    content['tables'].append(table_rows)
            
            return content
        
        except Exception as e:
            print(f"Error processing {doc_path}: {e}")
            return {'paragraphs': [], 'tables': [], 'topic_excerpts': {}, 'all_text': []}
    
    def categorize_text(self, text: str, topic_excerpts: Dict[str, List[str]]):
        """Categorize text by configured topics"""
        text_lower = text.lower()
        
        for topic, keywords in self.config.topic_mappings.items():
            if any(keyword.lower() in text_lower for keyword in keywords):
                if len(text) > self.config.excerpt_max_length:
                    text = text[:self.config.excerpt_max_length-3] + "..."
                topic_excerpts[topic].append(text)
    
    def generate_evidence_section(self, markdown_content: str) -> str:
        """Generate comprehensive evidence section"""
        if not self.source_contents:
            return markdown_content
        
        evidence_section = "\n\n## Supporting Evidence and Documentation\n\n"
        
        # Add excerpts from each source document
        for doc_path, content in self.source_contents.items():
            doc_name = Path(doc_path).stem
            evidence_section += f"### From {doc_name}\n\n"
            
            for topic, excerpts in content['topic_excerpts'].items():
                if excerpts:
                    topic_display = topic.replace('_', ' ').title()
                    evidence_section += f"**{topic_display}:**\n\n"
                    
                    # Limit excerpts per topic
                    limited_excerpts = excerpts[:self.config.max_excerpts_per_topic]
                    for i, excerpt in enumerate(limited_excerpts, 1):
                        evidence_section += f"{i}. *\"{excerpt}\"*\n\n"
        
        # Add cross-reference analysis if enabled
        if self.config.cross_reference_analysis and len(self.source_contents) >= 2:
            evidence_section += self.generate_cross_reference_analysis()
        
        return markdown_content + evidence_section
    
    def generate_cross_reference_analysis(self) -> str:
        """Generate cross-document analysis"""
        analysis_section = "### Cross-Document Analysis\n\n"
        
        # Find topics present in multiple documents
        topic_coverage = {}
        doc_names = [Path(doc_path).stem for doc_path in self.source_contents.keys()]
        
        for topic in self.config.topic_mappings.keys():
            coverage = []
            total_mentions = 0
            
            for doc_path, content in self.source_contents.items():
                doc_name = Path(doc_path).stem
                mentions = len(content['topic_excerpts'].get(topic, []))
                if mentions > 0:
                    coverage.append(f"{doc_name}: {mentions}")
                    total_mentions += mentions
            
            if len(coverage) >= 2:  # Topic appears in multiple documents
                topic_display = topic.replace('_', ' ').title()
                analysis_section += f"**{topic_display} Alignment:** "
                analysis_section += f"This topic appears across {len(coverage)} documents "
                analysis_section += f"({', '.join(coverage)} mentions), indicating strong institutional focus.\n\n"
        
        # Add thematic connections
        analysis_section += "**Key Thematic Connections:**\n\n"
        
        connection_count = 0
        for topic, keywords in self.config.topic_mappings.items():
            docs_with_topic = []
            sample_excerpts = []
            
            for doc_path, content in self.source_contents.items():
                if content['topic_excerpts'].get(topic, []):
                    docs_with_topic.append(Path(doc_path).stem)
                    if content['topic_excerpts'][topic]:
                        sample_excerpts.append(content['topic_excerpts'][topic][0])
            
            if len(docs_with_topic) >= 2:
                connection_count += 1
                topic_display = topic.replace('_', ' ').title()
                analysis_section += f"{connection_count}. **{topic_display}** bridges "
                analysis_section += f"{' and '.join(docs_with_topic)} through shared focus on "
                analysis_section += f"{', '.join(keywords[:3])}.\n\n"
        
        return analysis_section
    
    def enhance_markdown_file(self, markdown_path: str, output_path: Optional[str] = None) -> str:
        """Enhance a markdown file with evidence from source documents"""
        # Load source documents if not already loaded
        if not self.source_contents:
            self.load_source_documents()
        
        # Read markdown content
        with open(markdown_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Generate enhanced content
        enhanced_content = self.generate_evidence_section(content)
        
        # Save enhanced version
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(enhanced_content)
            print(f"✅ Enhanced version saved: {output_path}")
        
        return enhanced_content

def main():
    """Command line interface"""
    parser = argparse.ArgumentParser(description='Universal Document Enhancement Framework')
    parser.add_argument('markdown_file', help='Markdown file to enhance')
    parser.add_argument('--config', help='Configuration file (JSON)')
    parser.add_argument('--sources', nargs='+', help='Source documents (Word files)')
    parser.add_argument('--output', help='Output file path')
    parser.add_argument('--create-config', help='Create template configuration file')
    
    args = parser.parse_args()
    
    # Create template config if requested
    if args.create_config:
        config = EnhancementConfig()
        config.save_template(args.create_config)
        return
    
    # Load configuration
    if args.config:
        config = EnhancementConfig.from_json(args.config)
    else:
        # Use defaults with provided sources
        config = EnhancementConfig()
        if args.sources:
            config.source_documents = args.sources
        
        # Default topic mappings for LC documents
        config.topic_mappings = {
            "AI_Technology": ["AI", "artificial intelligence", "AI-assisted", "AI competency", 
                             "AI literacy", "AI training", "AI tools", "AI integration"],
            "Assessment": ["assessment", "evaluation", "rubric", "criteria", "testing", "quality assurance"],
            "Curriculum": ["curriculum", "OBTL", "PILO", "CILO", "learning outcomes", "course design"],
            "Multilingual": ["multilingual", "English", "Chinese", "language learning", "bilingual"],
            "Professional_Development": ["professional development", "training", "workshop", "capacity building"],
            "International": ["international", "exchange", "global citizenship", "intercultural"]
        }
    
    # Initialize enhancer
    enhancer = UniversalDocumentEnhancer(config)
    
    # Determine output path
    output_path = args.output or f"enhanced_{Path(args.markdown_file).name}"
    
    # Enhance markdown file
    print(f"🔧 Enhancing: {args.markdown_file}")
    enhancer.enhance_markdown_file(args.markdown_file, output_path)
    
    print("🎯 Enhancement complete!")

if __name__ == "__main__":
    main()
