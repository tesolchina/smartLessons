#!/usr/bin/env python3
"""
Generic Markdown Enhancement Tool
Enhances existing markdown files by adding excerpts from source documents
"""

import docx
import re
import json
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import argparse

class MarkdownEnhancer:
    """Generic tool for enhancing markdown files with document excerpts"""
    
    def __init__(self):
        self.supported_keywords = {
            'ai_related': ['AI', 'artificial intelligence', 'AI-assisted', 'AI-based', 
                          'AI competency', 'AI literacy', 'AI training', 'AI tools',
                          'AI integration', 'AI workshops', 'AI Task Force'],
            'assessment': ['assessment', 'evaluation', 'grading', 'rubric', 'criteria'],
            'curriculum': ['curriculum', 'course design', 'OBTL', 'PILO', 'CILO'],
            'development': ['professional development', 'training', 'workshop'],
            'international': ['international', 'exchange', 'global citizenship'],
            'community': ['community', 'public discourse', 'social impact']
        }
    
    def extract_word_document(self, doc_path: str) -> Dict[str, List[str]]:
        """Extract paragraphs and tables from Word document"""
        try:
            doc = docx.Document(doc_path)
            content = {
                'paragraphs': [],
                'tables': [],
                'all_text': []
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append(para.text.strip())
                    content['all_text'].append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_text.append(row_text)
                        content['all_text'].append(row_text)
                content['tables'].append(table_text)
            
            return content
            
        except Exception as e:
            print(f"Error extracting {doc_path}: {e}")
            return {'paragraphs': [], 'tables': [], 'all_text': []}
    
    def find_relevant_excerpts(self, content: Dict[str, List[str]], 
                              topic_keywords: List[str], 
                              max_excerpts: int = 3,
                              excerpt_length: int = 200) -> List[str]:
        """Find most relevant excerpts based on topic keywords"""
        relevant_excerpts = []
        
        for text in content['all_text']:
            # Count keyword matches
            matches = sum(1 for keyword in topic_keywords 
                         if keyword.lower() in text.lower())
            
            if matches > 0:
                # Create excerpt with context
                excerpt = self.create_excerpt(text, topic_keywords, excerpt_length)
                relevant_excerpts.append({
                    'text': excerpt,
                    'relevance_score': matches,
                    'length': len(text)
                })
        
        # Sort by relevance and return top excerpts
        relevant_excerpts.sort(key=lambda x: (x['relevance_score'], x['length']), 
                              reverse=True)
        
        return [excerpt['text'] for excerpt in relevant_excerpts[:max_excerpts]]
    
    def create_excerpt(self, text: str, keywords: List[str], max_length: int = 200) -> str:
        """Create focused excerpt around keywords"""
        # Find first keyword position
        keyword_pos = -1
        found_keyword = ""
        
        for keyword in keywords:
            pos = text.lower().find(keyword.lower())
            if pos != -1:
                keyword_pos = pos
                found_keyword = keyword
                break
        
        if keyword_pos == -1:
            return text[:max_length] + "..." if len(text) > max_length else text
        
        # Calculate excerpt bounds
        start = max(0, keyword_pos - max_length // 2)
        end = min(len(text), keyword_pos + max_length // 2)
        
        excerpt = text[start:end]
        
        # Add ellipsis if truncated
        if start > 0:
            excerpt = "..." + excerpt
        if end < len(text):
            excerpt = excerpt + "..."
        
        return excerpt
    
    def enhance_markdown_section(self, markdown_content: str, 
                                section_title: str,
                                excerpts: List[str],
                                source_doc: str) -> str:
        """Enhance a specific section with excerpts"""
        
        # Find the section
        section_pattern = rf"(#{1,6}\s*{re.escape(section_title)}.*?)(?=#{1,6}|\Z)"
        match = re.search(section_pattern, markdown_content, re.DOTALL | re.IGNORECASE)
        
        if not match:
            print(f"Section '{section_title}' not found")
            return markdown_content
        
        section_content = match.group(1)
        
        # Add excerpts if not already present
        if "### Supporting Evidence" not in section_content:
            excerpts_text = f"\n\n### Supporting Evidence from {Path(source_doc).name}:\n\n"
            
            for i, excerpt in enumerate(excerpts, 1):
                excerpts_text += f"{i}. *\"{excerpt}\"*\n\n"
            
            enhanced_section = section_content + excerpts_text
            
            return markdown_content.replace(section_content, enhanced_section)
        
        return markdown_content
    
    def enhance_markdown_file(self, markdown_path: str, 
                             source_documents: List[str],
                             topics_mapping: Dict[str, List[str]]) -> str:
        """Enhance entire markdown file with excerpts from source documents"""
        
        # Read existing markdown
        with open(markdown_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        enhanced_content = markdown_content
        
        # Process each source document
        for doc_path in source_documents:
            print(f"📄 Processing excerpts from: {Path(doc_path).name}")
            doc_content = self.extract_word_document(doc_path)
            
            # Enhance sections based on topics mapping
            for section_title, keywords in topics_mapping.items():
                excerpts = self.find_relevant_excerpts(doc_content, keywords)
                if excerpts:
                    enhanced_content = self.enhance_markdown_section(
                        enhanced_content, section_title, excerpts, doc_path
                    )
        
        return enhanced_content
    
    def auto_detect_topics(self, markdown_content: str) -> Dict[str, List[str]]:
        """Auto-detect topics from markdown headings and suggest keywords"""
        topics = {}
        
        # Find all headings
        heading_pattern = r"#{1,6}\s+(.+)"
        headings = re.findall(heading_pattern, markdown_content)
        
        for heading in headings:
            heading_clean = heading.strip()
            suggested_keywords = []
            
            # Map headings to keyword categories
            heading_lower = heading_clean.lower()
            
            if any(term in heading_lower for term in ['ai', 'artificial', 'technology']):
                suggested_keywords.extend(self.supported_keywords['ai_related'])
            
            if any(term in heading_lower for term in ['assessment', 'evaluation', 'quality']):
                suggested_keywords.extend(self.supported_keywords['assessment'])
            
            if any(term in heading_lower for term in ['curriculum', 'teaching', 'learning']):
                suggested_keywords.extend(self.supported_keywords['curriculum'])
            
            if any(term in heading_lower for term in ['development', 'training', 'staff']):
                suggested_keywords.extend(self.supported_keywords['development'])
            
            if any(term in heading_lower for term in ['international', 'global', 'exchange']):
                suggested_keywords.extend(self.supported_keywords['international'])
            
            if any(term in heading_lower for term in ['community', 'public', 'social']):
                suggested_keywords.extend(self.supported_keywords['community'])
            
            if suggested_keywords:
                topics[heading_clean] = list(set(suggested_keywords))
        
        return topics

def main():
    """Main CLI interface"""
    parser = argparse.ArgumentParser(description='Enhance Markdown files with document excerpts')
    parser.add_argument('markdown_file', help='Markdown file to enhance')
    parser.add_argument('source_docs', nargs='+', help='Source Word documents')
    parser.add_argument('--output', help='Output file (default: enhanced_[original].md)')
    parser.add_argument('--topics', help='JSON file with topics mapping')
    parser.add_argument('--auto-detect', action='store_true', 
                       help='Auto-detect topics from markdown headings')
    
    args = parser.parse_args()
    
    enhancer = MarkdownEnhancer()
    
    # Determine topics mapping
    topics_mapping = {}
    
    if args.topics:
        with open(args.topics, 'r', encoding='utf-8') as f:
            topics_mapping = json.load(f)
    elif args.auto_detect:
        with open(args.markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        topics_mapping = enhancer.auto_detect_topics(markdown_content)
        print("🎯 Auto-detected topics:", list(topics_mapping.keys()))
    else:
        # Default mapping for LC documents
        topics_mapping = {
            "AI": enhancer.supported_keywords['ai_related'],
            "Assessment": enhancer.supported_keywords['assessment'],
            "Curriculum": enhancer.supported_keywords['curriculum'],
            "Development": enhancer.supported_keywords['development']
        }
    
    # Enhance markdown file
    enhanced_content = enhancer.enhance_markdown_file(
        args.markdown_file, args.source_docs, topics_mapping
    )
    
    # Save enhanced file
    output_file = args.output or f"enhanced_{Path(args.markdown_file).name}"
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(enhanced_content)
    
    print(f"✅ Enhanced markdown saved to: {output_file}")

if __name__ == "__main__":
    main()
