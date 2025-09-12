#!/usr/bin/env python3
"""
Word Document Edit Locator for Innovation Officer
Identifies specific paragraphs and sections in the original Word document for afternoon edits
"""

import docx
from pathlib import Path
from typing import Dict, List, Tuple

class WordDocEditLocator:
    """Locate specific edits in the original Word document"""
    
    def __init__(self, doc_path: str):
        self.doc_path = doc_path
        self.doc = docx.Document(doc_path)
        self.paragraphs = [para for para in self.doc.paragraphs if para.text.strip()]
        
    def find_ai_enhancement_targets(self) -> List[Dict]:
        """Find paragraphs that mention AI and need enhancement"""
        targets = []
        
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            
            # Look for AI competency mentions that need strengthening
            if 'AI competency' in text and 'measurable' not in text:
                targets.append({
                    'paragraph_number': i + 1,
                    'section_context': self.get_section_context(i),
                    'current_text': text[:100] + "..." if len(text) > 100 else text,
                    'full_text': text,
                    'edit_type': 'ENHANCE_AI_COMPETENCY',
                    'suggested_change': text.replace('AI competency', 'measurable AI competency with rubric-based assessment'),
                    'reason': 'Makes AI competency concrete and assessable',
                    'priority': 'HIGH',
                    'time_estimate': '1 minute'
                })
            
            # Look for vague implementation language
            if ('will be implemented' in text or 'will be integrated' in text) and 'by' not in text and len(text) > 50:
                enhanced_text = text.replace('will be implemented', 'will be implemented by AY 2025-26')
                enhanced_text = enhanced_text.replace('will be integrated', 'will be integrated by AY 2025-26')
                
                targets.append({
                    'paragraph_number': i + 1,
                    'section_context': self.get_section_context(i),
                    'current_text': text[:100] + "..." if len(text) > 100 else text,
                    'full_text': text,
                    'edit_type': 'ADD_TIMELINE',
                    'suggested_change': enhanced_text,
                    'reason': 'Adds specific timeline for accountability',
                    'priority': 'MEDIUM',
                    'time_estimate': '30 seconds'
                })
            
            # Look for training mentions without specificity
            if 'training' in text.lower() and 'AI' in text and 'workshop' not in text and len(text) > 30:
                enhanced_text = text.replace('training', 'monthly AI innovation workshops and training')
                
                targets.append({
                    'paragraph_number': i + 1,
                    'section_context': self.get_section_context(i),
                    'current_text': text[:100] + "..." if len(text) > 100 else text,
                    'full_text': text,
                    'edit_type': 'SPECIFY_TRAINING',
                    'suggested_change': enhanced_text,
                    'reason': 'Specifies regular AI training schedule',
                    'priority': 'HIGH',
                    'time_estimate': '45 seconds'
                })
        
        return targets
    
    def find_strategic_insertion_points(self) -> List[Dict]:
        """Find places to insert Innovation Officer initiatives"""
        insertion_points = []
        
        for i, para in enumerate(self.paragraphs):
            text = para.text.strip()
            
            # After strategic direction sections
            if ('strategic direction' in text.lower() or 'mission' in text.lower()) and len(text) > 50:
                insertion_points.append({
                    'insert_after_paragraph': i + 1,
                    'section_context': self.get_section_context(i),
                    'anchor_text': text[:80] + "..." if len(text) > 80 else text,
                    'text_to_insert': 'AI Innovation Priority: The Language Center will integrate measurable AI competency development across all credit-bearing courses by AY 2025-26, with rubric-based assessment and faculty training support.',
                    'reason': 'Establishes concrete Innovation Officer leadership initiative',
                    'priority': 'CRITICAL',
                    'time_estimate': '1 minute'
                })
            
            # After assessment sections
            if 'assessment' in text.lower() and 'quality' in text.lower() and len(text) > 40:
                insertion_points.append({
                    'insert_after_paragraph': i + 1,
                    'section_context': self.get_section_context(i),
                    'anchor_text': text[:80] + "..." if len(text) > 80 else text,
                    'text_to_insert': 'AI Assessment Innovation: 50% of LC courses will pilot AI-enhanced feedback systems and automated rubric application by end of AY 2025-26, with effectiveness measured through student learning outcomes and faculty feedback.',
                    'reason': 'Concrete AI assessment initiative with measurable targets',
                    'priority': 'HIGH',
                    'time_estimate': '1 minute'
                })
            
            # After professional development sections
            if ('professional development' in text.lower() or 'faculty development' in text.lower()) and len(text) > 30:
                insertion_points.append({
                    'insert_after_paragraph': i + 1,
                    'section_context': self.get_section_context(i),
                    'anchor_text': text[:80] + "..." if len(text) > 80 else text,
                    'text_to_insert': 'AI Faculty Development Framework: Monthly AI Innovation Workshops covering (1) AI tool evaluation and pedagogical integration, (2) Ethical AI usage in language education, (3) AI-assisted assessment methodologies, (4) Student AI literacy development. Attendance tracking and competency certification required.',
                    'reason': 'Shows systematic approach to building AI capacity',
                    'priority': 'HIGH',
                    'time_estimate': '1.5 minutes'
                })
        
        return insertion_points
    
    def get_section_context(self, para_index: int) -> str:
        """Get the section heading context for a paragraph"""
        # Look backwards for the most recent heading-style paragraph
        for i in range(para_index, -1, -1):
            para = self.paragraphs[i]
            if para.style and 'Heading' in para.style.name:
                return f"Under section: {para.text.strip()}"
            # Also check for bold or short paragraphs that might be headings
            if len(para.text.strip()) < 50 and para.text.strip().isupper():
                return f"Under section: {para.text.strip()}"
        
        return f"Around paragraph {para_index + 1}"
    
    def create_word_edit_guide(self) -> str:
        """Create afternoon edit guide for Word document"""
        
        enhancement_targets = self.find_ai_enhancement_targets()
        insertion_points = self.find_strategic_insertion_points()
        
        # Sort by priority
        high_priority_edits = [t for t in enhancement_targets if t['priority'] == 'HIGH']
        medium_priority_edits = [t for t in enhancement_targets if t['priority'] == 'MEDIUM']
        
        critical_insertions = [i for i in insertion_points if i['priority'] == 'CRITICAL']
        high_priority_insertions = [i for i in insertion_points if i['priority'] == 'HIGH']
        
        total_high_priority = len(high_priority_edits) + len(critical_insertions)
        total_items = len(enhancement_targets) + len(insertion_points)
        
        guide = f"""# Word Document Edit Session for Innovation Officer
**Document:** LC DAA Report_response_20251009.docx
**Total High-Priority Items:** {total_high_priority}
**Estimated Time:** 2-3 hours

## Phase 1: Critical & High Priority (45-60 minutes)

"""
        
        # Add critical insertions first
        for i, item in enumerate(critical_insertions, 1):
            guide += f"""### {i}. CRITICAL INSERTION 🔴
**Location:** After paragraph {item['insert_after_paragraph']}
**Context:** {item['section_context']}
**Anchor text:** "{item['anchor_text']}"

**INSERT THIS TEXT:**
> {item['text_to_insert']}

**Why:** {item['reason']}
**Time needed:** {item['time_estimate']}

---

"""
        
        # Add high priority edits
        for i, item in enumerate(high_priority_edits, len(critical_insertions) + 1):
            guide += f"""### {i}. HIGH PRIORITY EDIT 🟡
**Location:** Paragraph {item['paragraph_number']}
**Context:** {item['section_context']}
**Current text:** "{item['current_text']}"

**CHANGE TO:**
> {item['suggested_change']}

**Why:** {item['reason']}
**Time needed:** {item['time_estimate']}

---

"""
        
        # Add high priority insertions
        for i, item in enumerate(high_priority_insertions, len(critical_insertions) + len(high_priority_edits) + 1):
            guide += f"""### {i}. HIGH PRIORITY INSERTION 🟡
**Location:** After paragraph {item['insert_after_paragraph']}
**Context:** {item['section_context']}
**Anchor text:** "{item['anchor_text']}"

**INSERT THIS TEXT:**
> {item['text_to_insert']}

**Why:** {item['reason']}
**Time needed:** {item['time_estimate']}

---

"""
        
        guide += "## Phase 2: Medium Priority (30-45 minutes)\n\n"
        
        for i, item in enumerate(medium_priority_edits, 1):
            guide += f"""### {i}. Medium Priority Edit
**Paragraph {item['paragraph_number']}** - {item['reason']}
**Change:** "{item['current_text'][:50]}..." → Enhanced with timeline

"""
        
        guide += f"""
## Quick Reference for Word Document Navigation

1. **Open:** LC DAA Report_response_20251009.docx
2. **Find paragraphs:** Use Ctrl+F (Cmd+F) to search for anchor text
3. **Edit efficiently:** 
   - For edits: Select text → Replace with provided text
   - For insertions: Place cursor at end of anchor paragraph → Press Enter → Add new text
4. **Save regularly:** Save after every 3-4 edits

## Success Checklist
- [ ] {len(critical_insertions)} CRITICAL insertions completed
- [ ] {len(high_priority_edits)} HIGH priority edits completed  
- [ ] {len(high_priority_insertions)} HIGH priority insertions completed
- [ ] Document saved with Innovation Officer enhancements
- [ ] AI competency mentions made measurable and concrete

## Time Breakdown
- **Critical items:** {len(critical_insertions)} × 1-1.5 min = {len(critical_insertions) * 1.25:.0f} minutes
- **High priority edits:** {len(high_priority_edits)} × 0.75 min = {len(high_priority_edits) * 0.75:.0f} minutes  
- **High priority insertions:** {len(high_priority_insertions)} × 1 min = {len(high_priority_insertions):.0f} minutes
- **Total Phase 1:** {(len(critical_insertions) * 1.25 + len(high_priority_edits) * 0.75 + len(high_priority_insertions)):.0f} minutes

*All edits focus on establishing Innovation Officer leadership and making AI initiatives concrete and measurable.*
"""
        
        return guide
    
    def save_word_edit_guide(self):
        """Generate and save Word document edit guide"""
        
        print("🔍 Analyzing Word document for Innovation Officer enhancements...")
        guide_content = self.create_word_edit_guide()
        
        # Save guide
        output_dir = Path(self.doc_path).parent
        guide_path = output_dir / 'Word_Document_Edit_Session_Guide.md'
        
        with open(guide_path, 'w', encoding='utf-8') as f:
            f.write(guide_content)
        
        print(f"✅ Word document edit guide saved: {guide_path}")
        
        # Print summary
        enhancement_targets = self.find_ai_enhancement_targets()
        insertion_points = self.find_strategic_insertion_points()
        
        high_priority_count = len([t for t in enhancement_targets if t['priority'] == 'HIGH'])
        critical_count = len([i for i in insertion_points if i['priority'] == 'CRITICAL'])
        
        print(f"\n📊 Edit Session Summary:")
        print(f"   • {critical_count} CRITICAL insertions")
        print(f"   • {high_priority_count} HIGH priority edits")
        print(f"   • Total items: {len(enhancement_targets) + len(insertion_points)}")
        print(f"   • Focus: Making AI initiatives concrete and measurable")

def main():
    """Generate Word document edit session"""
    
    doc_path = "/Users/simonwang/Documents/Usage/VibeCoding/DailyAssistant/projects/LC_admin/PMC/LC DAA Report_response_20251009.docx"
    
    print("📝 Creating Word document edit session for Innovation Officer...")
    
    try:
        locator = WordDocEditLocator(doc_path)
        locator.save_word_edit_guide()
        
        print("\n🎯 Next Steps:")
        print("1. Open 'Word_Document_Edit_Session_Guide.md'")
        print("2. Open the Word document: LC DAA Report_response_20251009.docx")
        print("3. Follow paragraph-by-paragraph edit instructions")
        print("4. Start with CRITICAL insertions, then HIGH priority edits")
        print("5. Use Ctrl+F (Cmd+F) to find anchor texts quickly")
        
    except Exception as e:
        print(f"❌ Error processing Word document: {e}")
        print("Please check that the document exists and is accessible.")

if __name__ == "__main__":
    main()
