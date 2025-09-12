import os
from docx import Document

def add_url_to_docs(folder_path):
    """
    Add the corresponding URL to the top of each Word document in the folder.
    """
    try:
        # Loop through all files in the folder and its subfolders
        for root, _, files in os.walk(folder_path):
            for file in files:
                if file.endswith(".docx"):
                    file_path = os.path.join(root, file)

                    # Determine the URL from the folder structure (replace this with actual logic if needed)
                    relative_path = os.path.relpath(file_path, folder_path)
                    url = "https://lc.hkbu.edu.hk/main/" + "/".join(relative_path.split(os.sep)[:-1])

                    # Open the Word document
                    doc = Document(file_path)

                    # Add the URL at the top of the document
                    doc.paragraphs.insert(0, doc.add_paragraph(f"URL: {url}\n"))

                    # Save the updated document
                    doc.save(file_path)
                    print(f"Updated: {file_path} with URL: {url}")
    except Exception as e:
        print(f"Error processing Word files: {e}")

if __name__ == "__main__":
    # Path to the folder containing Word documents
    folder_path = "/workspaces/lcwebsiteAnalytics/.vscode/results"

    # Add URLs to Word documents
    add_url_to_docs(folder_path)